"""ÎßàÌÅ¨Îã§Ïö¥ÏùÑ Word Î¨∏ÏÑúÎ°ú Î≥ÄÌôò (Google Docs ÏµúÏ†ÅÌôî)"""
import os
import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Í≤ΩÎ°ú ÏÑ§Ï†ï
README_PATH = r'D:\AI\claude01\wsoptv_v2\tasks\prds\0010-prd-wsoptv\README.md'
OUTPUT_DIR = r'D:\AI\claude01\wsoptv_v2\tasks\prds\0010-prd-wsoptv'
TEMP_DIR = tempfile.mkdtemp()

print(f"ÏûÑÏãú ÎîîÎ†âÌÜ†Î¶¨: {TEMP_DIR}")

# README.md ÏùΩÍ∏∞
with open(README_PATH, 'r', encoding='utf-8') as f:
    content = f.read()

# Î≤ÑÏ†Ñ Ï∂îÏ∂ú
version_match = re.search(r'\*\*Version\*\*:\s*(\d+\.\d+\.\d+)', content)
version = version_match.group(1) if version_match else datetime.now().strftime('%Y%m%d')

OUTPUT_PATH = os.path.join(OUTPUT_DIR, f'WSOPTV_PRD_v{version}.docx')
print(f"Î≤ÑÏ†Ñ: {version}")
print(f"Ï∂úÎ†• ÌååÏùº: {OUTPUT_PATH}")

# ============================================
# Mermaid Ïù¥ÎØ∏ÏßÄ Î≥ÄÌôò
# ============================================
mermaid_pattern = r'```mermaid\n(.*?)```'
mermaid_blocks = re.findall(mermaid_pattern, content, re.DOTALL)

print(f"Î∞úÍ≤¨Îêú Mermaid Îã§Ïù¥Ïñ¥Í∑∏Îû®: {len(mermaid_blocks)}Í∞ú")

mermaid_images = {}
for i, mermaid_code in enumerate(mermaid_blocks):
    mmd_file = os.path.join(TEMP_DIR, f'diagram_{i}.mmd')
    png_file = os.path.join(TEMP_DIR, f'diagram_{i}.png')

    with open(mmd_file, 'w', encoding='utf-8') as f:
        f.write(mermaid_code.strip())

    try:
        result = subprocess.run(
            ['mmdc', '-i', mmd_file, '-o', png_file, '-b', 'white', '-s', '2'],
            capture_output=True,
            text=True,
            timeout=30,
            shell=True
        )
        if os.path.exists(png_file):
            mermaid_images[i] = png_file
            print(f"  [{i+1}/{len(mermaid_blocks)}] Î≥ÄÌôò ÏôÑÎ£å")
        else:
            print(f"  [{i+1}/{len(mermaid_blocks)}] Î≥ÄÌôò Ïã§Ìå®")
    except Exception as e:
        print(f"  [{i+1}/{len(mermaid_blocks)}] Ïò§Î•ò: {e}")

print(f"\nÏù¥ÎØ∏ÏßÄ Î≥ÄÌôò ÏôÑÎ£å: {len(mermaid_images)}Í∞ú")

# ============================================
# Ìó¨Ìçº Ìï®Ïàò
# ============================================
def set_cell_shading(cell, color):
    """ÏÖÄ Î∞∞Í≤ΩÏÉâ ÏÑ§Ï†ï"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def parse_markdown_table(lines):
    """ÎßàÌÅ¨Îã§Ïö¥ ÌÖåÏù¥Î∏î ÌååÏã±"""
    rows = []
    for line in lines:
        if line.startswith('|') and not line.startswith('|--') and not line.startswith('|-'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and not all(c.replace('-', '').replace(':', '') == '' for c in cells):
                rows.append(cells)
    return rows

def clean_text(text):
    """ÎßàÌÅ¨Îã§Ïö¥ ÏÑúÏãù Ï†úÍ±∞ Î∞è Ï†ïÎ¶¨"""
    # Î≥ºÎìú Ï≤òÎ¶¨
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Ïù¥ÌÉ§Î¶≠ Ï≤òÎ¶¨
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # ÎßÅÌÅ¨ Ï≤òÎ¶¨ [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Ïù∏ÎùºÏù∏ ÏΩîÎìú
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()

def add_formatted_paragraph(doc, text, bold_patterns=None):
    """Î≥ºÎìú ÏÑúÏãùÏù¥ Ï†ÅÏö©Îêú Îã®ÎùΩ Ï∂îÍ∞Ä"""
    p = doc.add_paragraph()

    # **text** Ìå®ÌÑ¥ Ï∞æÍ∏∞
    pattern = r'\*\*(.+?)\*\*'
    last_end = 0

    for match in re.finditer(pattern, text):
        # Î≥ºÎìú Ï†Ñ ÏùºÎ∞ò ÌÖçÏä§Ìä∏
        if match.start() > last_end:
            p.add_run(text[last_end:match.start()])
        # Î≥ºÎìú ÌÖçÏä§Ìä∏
        bold_run = p.add_run(match.group(1))
        bold_run.bold = True
        last_end = match.end()

    # ÎÇ®ÏùÄ ÌÖçÏä§Ìä∏
    if last_end < len(text):
        p.add_run(text[last_end:])

    return p

# ============================================
# Word Î¨∏ÏÑú ÏÉùÏÑ±
# ============================================
doc = Document()

# Í∏∞Î≥∏ Ïä§ÌÉÄÏùº ÏÑ§Ï†ï
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(11)

# Ï†úÎ™© Ïä§ÌÉÄÏùº
for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Malgun Gothic'
    heading_style.font.color.rgb = RGBColor(0, 51, 102)

# ============================================
# ÎßàÌÅ¨Îã§Ïö¥ ÌååÏã± Î∞è Î≥ÄÌôò
# ============================================
lines = content.split('\n')
i = 0
mermaid_index = 0
table_buffer = []
in_table = False

while i < len(lines):
    line = lines[i]

    # ========== Mermaid ÏΩîÎìú Î∏îÎ°ù ==========
    if line.startswith('```mermaid'):
        # ÎÅùÍπåÏßÄ Ïä§ÌÇµ
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            i += 1

        # Ïù¥ÎØ∏ÏßÄ ÏÇΩÏûÖ
        if mermaid_index in mermaid_images:
            try:
                doc.add_picture(mermaid_images[mermaid_index], width=Inches(5.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Îã§Ïù¥Ïñ¥Í∑∏Îû® {mermaid_index + 1}]")

        mermaid_index += 1
        i += 1
        continue

    # ========== ÏùºÎ∞ò ÏΩîÎìú Î∏îÎ°ù ==========
    if line.startswith('```'):
        code_lang = line[3:].strip()
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1

        if code_lines:
            # ÏΩîÎìú Î∏îÎ°ù Ïä§ÌÉÄÏùº
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            # Î∞∞Í≤ΩÏÉâ Ìö®Í≥ºÎ•º ÏúÑÌïú Îì§Ïó¨Ïì∞Í∏∞
            p.paragraph_format.left_indent = Cm(0.5)

        i += 1
        continue

    # ========== ÌÖåÏù¥Î∏î Ï≤òÎ¶¨ ==========
    if line.startswith('|'):
        if not in_table:
            in_table = True
            table_buffer = []
        table_buffer.append(line)
        i += 1
        continue
    elif in_table:
        # ÌÖåÏù¥Î∏î Ï¢ÖÎ£å, ÌÖåÏù¥Î∏î ÏÉùÏÑ±
        in_table = False
        rows = parse_markdown_table(table_buffer)

        if rows and len(rows) > 0:
            num_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = clean_text(cell_text)

                        # Ìó§Îçî Ìñâ Ïä§ÌÉÄÏùº
                        if row_idx == 0:
                            set_cell_shading(cell, '003366')
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)

                        # ÏÖÄ Ï†ïÎ†¨
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ÌÖåÏù¥Î∏î ÌõÑ Îπà Ï§Ñ
            doc.add_paragraph()

        table_buffer = []

    # ========== Ï†úÎ™© ==========
    if line.startswith('# ') and not line.startswith('##'):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue
    elif line.startswith('## ') and not line.startswith('###'):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    elif line.startswith('### ') and not line.startswith('####'):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue
    elif line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=4)
        i += 1
        continue

    # ========== Ïù∏Ïö©Î¨∏ ==========
    if line.startswith('> '):
        quote_text = line[2:].strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)

        # Î≥ºÎìú Ï≤òÎ¶¨
        pattern = r'\*\*(.+?)\*\*'
        last_end = 0
        for match in re.finditer(pattern, quote_text):
            if match.start() > last_end:
                run = p.add_run(quote_text[last_end:match.start()])
                run.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
            bold_run = p.add_run(match.group(1))
            bold_run.bold = True
            bold_run.italic = True
            bold_run.font.color.rgb = RGBColor(80, 80, 80)
            last_end = match.end()
        if last_end < len(quote_text):
            run = p.add_run(quote_text[last_end:])
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)

        i += 1
        continue

    # ========== Î™©Î°ù ==========
    if line.startswith('- '):
        list_text = line[2:].strip()
        p = doc.add_paragraph(style='List Bullet')

        # Î≥ºÎìú Ï≤òÎ¶¨
        pattern = r'\*\*(.+?)\*\*'
        last_end = 0
        clean_list_text = list_text
        for match in re.finditer(pattern, list_text):
            if match.start() > last_end:
                p.add_run(list_text[last_end:match.start()])
            bold_run = p.add_run(match.group(1))
            bold_run.bold = True
            last_end = match.end()
        if last_end < len(list_text):
            p.add_run(list_text[last_end:])
        elif last_end == 0:
            p.add_run(list_text)

        i += 1
        continue

    # ========== Íµ¨Î∂ÑÏÑ† ==========
    if line.startswith('---'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('‚îÅ' * 40)
        run.font.color.rgb = RGBColor(200, 200, 200)
        i += 1
        continue

    # ========== ÏùºÎ∞ò ÌÖçÏä§Ìä∏ ==========
    if line.strip():
        add_formatted_paragraph(doc, line)

    i += 1

# ÎßàÏßÄÎßâ ÌÖåÏù¥Î∏î Ï≤òÎ¶¨
if in_table and table_buffer:
    rows = parse_markdown_table(table_buffer)
    if rows and len(rows) > 0:
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = clean_text(cell_text)
                    if row_idx == 0:
                        set_cell_shading(cell, '003366')
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)

# Ï†ÄÏû•
doc.save(OUTPUT_PATH)
print(f'\n‚úÖ Word Î¨∏ÏÑú ÏÉùÏÑ± ÏôÑÎ£å!')
print(f'üìÑ ÌååÏùº: {OUTPUT_PATH}')
print(f'\nÏù¥ ÌååÏùºÏùÑ Google DriveÏóê ÏóÖÎ°úÎìú ÌõÑ "Google Î¨∏ÏÑúÎ°ú Ïó¥Í∏∞"ÌïòÏÑ∏Ïöî.')

# ÏûÑÏãú ÌååÏùº Ï†ïÎ¶¨
import shutil
shutil.rmtree(TEMP_DIR, ignore_errors=True)
