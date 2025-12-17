"""마크다운을 Word 문서로 변환"""
from docx import Document
from docx.shared import Pt
import re

# README.md 읽기
with open(r'D:\AI\claude01\wsoptv_v2\tasks\prds\0010-prd-wsoptv\README.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Word 문서 생성
doc = Document()

# 마크다운을 간단히 파싱하여 Word로 변환
lines = content.split('\n')
in_code_block = False
code_buffer = []

for line in lines:
    # 코드 블록 처리
    if line.startswith('```'):
        if in_code_block:
            # 코드 블록 끝
            code_text = '\n'.join(code_buffer)
            if code_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            code_buffer = []
        in_code_block = not in_code_block
        continue

    if in_code_block:
        code_buffer.append(line)
        continue

    # 제목 처리
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('> '):
        # 인용문
        p = doc.add_paragraph(line[2:])
        p.style = 'Quote'
    elif line.startswith('| '):
        # 테이블 (간단히 텍스트로)
        doc.add_paragraph(line)
    elif line.startswith('- '):
        # 목록
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('---'):
        # 구분선
        doc.add_paragraph('─' * 50)
    elif line.strip():
        # 일반 텍스트
        doc.add_paragraph(line)

# 저장
output_path = r'D:\AI\claude01\wsoptv_v2\tasks\prds\0010-prd-wsoptv\WSOPTV_PRD.docx'
doc.save(output_path)
print(f'Word 문서 생성 완료!')
print(f'')
print(f'파일 경로: {output_path}')
print(f'')
print(f'이 파일을 Google Drive에 업로드하면 자동으로 Google Docs로 변환됩니다.')
