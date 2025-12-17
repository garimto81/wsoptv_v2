"""02-user-experience.md를 Word 문서로 변환 (Google Docs 최적화)"""
import os
import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 경로 설정
INPUT_PATH = r'D:\AI\claude01\wsoptv_v2\tasks\prds\0010-prd-wsoptv\02-user-experience.md'
OUTPUT_DIR = r'D:\AI\claude01\wsoptv_v2\tasks\prds\0010-prd-wsoptv'
TEMP_DIR = tempfile.mkdtemp()

print(f"입력 파일: {INPUT_PATH}")
print(f"임시 디렉토리: {TEMP_DIR}")

# 마크다운 읽기
with open(INPUT_PATH, 'r', encoding='utf-8') as f:
    content = f.read()

# 버전 추출
version_match = re.search(r'\*\*Version\*\*:\s*(\d+\.\d+\.\d+)', content)
version = version_match.group(1) if version_match else datetime.now().strftime('%Y%m%d')

OUTPUT_PATH = os.path.join(OUTPUT_DIR, f'WSOPTV_UX_v{version}.docx')
print(f"버전: {version}")
print(f"출력 파일: {OUTPUT_PATH}")

# ============================================
# Mermaid 이미지 변환
# ============================================
mermaid_pattern = r'```mermaid\n(.*?)```'
mermaid_blocks = re.findall(mermaid_pattern, content, re.DOTALL)

print(f"\n발견된 Mermaid 다이어그램: {len(mermaid_blocks)}개")

mermaid_images = {}
for i, mermaid_code in enumerate(mermaid_blocks):
    mmd_file = os.path.join(TEMP_DIR, f'diagram_{i}.mmd')
    png_file = os.path.join(TEMP_DIR, f'diagram_{i}.png')

    with open(mmd_file, 'w', encoding='utf-8') as f:
        f.write(mermaid_code.strip())

    try:
        result = subprocess.run(
            ['mmdc', '-i', mmd_file, '-o', png_file, '-b', 'white', '-s', '2'],
            capture_output=True,
            timeout=30,
            shell=True
        )
        if os.path.exists(png_file):
            mermaid_images[i] = png_file
            print(f"  [{i+1}/{len(mermaid_blocks)}] PNG 변환 완료")
        else:
            print(f"  [{i+1}/{len(mermaid_blocks)}] 변환 실패")
    except Exception as e:
        print(f"  [{i+1}/{len(mermaid_blocks)}] 오류: {e}")

print(f"\n이미지 변환 완료: {len(mermaid_images)}개")

# ============================================
# 헬퍼 함수
# ============================================
def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def parse_markdown_table(lines):
    """마크다운 테이블 파싱"""
    rows = []
    for line in lines:
        if line.startswith('|') and not line.startswith('|--') and not line.startswith('|-'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and not all(c.replace('-', '').replace(':', '') == '' for c in cells):
                rows.append(cells)
    return rows

def clean_text(text):
    """마크다운 서식 제거"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()

def add_formatted_paragraph(doc, text):
    """볼드 서식 적용 단락 추가"""
    p = doc.add_paragraph()
    pattern = r'\*\*(.+?)\*\*'
    last_end = 0

    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            p.add_run(text[last_end:match.start()])
        bold_run = p.add_run(match.group(1))
        bold_run.bold = True
        last_end = match.end()

    if last_end < len(text):
        p.add_run(text[last_end:])

    return p

# ============================================
# Word 문서 생성
# ============================================
doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(11)

# 제목 스타일
for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Malgun Gothic'
    heading_style.font.color.rgb = RGBColor(0, 51, 102)

# ============================================
# 마크다운 파싱 및 변환
# ============================================
lines = content.split('\n')
i = 0
mermaid_index = 0
table_buffer = []
in_table = False

while i < len(lines):
    line = lines[i]

    # Mermaid 코드 블록
    if line.startswith('```mermaid'):
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            i += 1

        if mermaid_index in mermaid_images:
            try:
                doc.add_picture(mermaid_images[mermaid_index], width=Inches(5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # 여백
            except Exception as e:
                doc.add_paragraph(f"[다이어그램 {mermaid_index + 1}]")

        mermaid_index += 1
        i += 1
        continue

    # 일반 코드 블록
    if line.startswith('```'):
        code_lang = line[3:].strip()
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1

        if code_lines:
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Cm(0.5)

        i += 1
        continue

    # 테이블 처리
    if line.startswith('|'):
        if not in_table:
            in_table = True
            table_buffer = []
        table_buffer.append(line)
        i += 1
        continue
    elif in_table:
        in_table = False
        rows = parse_markdown_table(table_buffer)

        if rows and len(rows) > 0:
            num_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = clean_text(cell_text)

                        # 헤더 행 스타일
                        if row_idx == 0:
                            set_cell_shading(cell, '003366')
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.font.size = Pt(10)

                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        table_buffer = []

    # 제목
    if line.startswith('# ') and not line.startswith('##'):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue
    elif line.startswith('## ') and not line.startswith('###'):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    elif line.startswith('### ') and not line.startswith('####'):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue
    elif line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=4)
        i += 1
        continue

    # 인용문
    if line.startswith('> '):
        quote_text = line[2:].strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)

        pattern = r'\*\*(.+?)\*\*'
        last_end = 0
        for match in re.finditer(pattern, quote_text):
            if match.start() > last_end:
                run = p.add_run(quote_text[last_end:match.start()])
                run.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
            bold_run = p.add_run(match.group(1))
            bold_run.bold = True
            bold_run.italic = True
            bold_run.font.color.rgb = RGBColor(80, 80, 80)
            last_end = match.end()
        if last_end < len(quote_text):
            run = p.add_run(quote_text[last_end:])
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)

        i += 1
        continue

    # 목록
    if line.startswith('- '):
        list_text = line[2:].strip()
        p = doc.add_paragraph(style='List Bullet')

        pattern = r'\*\*(.+?)\*\*'
        last_end = 0
        for match in re.finditer(pattern, list_text):
            if match.start() > last_end:
                p.add_run(list_text[last_end:match.start()])
            bold_run = p.add_run(match.group(1))
            bold_run.bold = True
            last_end = match.end()
        if last_end < len(list_text):
            p.add_run(list_text[last_end:])
        elif last_end == 0:
            p.add_run(list_text)

        i += 1
        continue

    # 구분선
    if line.startswith('---'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('_' * 50)
        run.font.color.rgb = RGBColor(200, 200, 200)
        i += 1
        continue

    # 일반 텍스트
    if line.strip():
        add_formatted_paragraph(doc, line)

    i += 1

# 마지막 테이블 처리
if in_table and table_buffer:
    rows = parse_markdown_table(table_buffer)
    if rows and len(rows) > 0:
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = clean_text(cell_text)
                    if row_idx == 0:
                        set_cell_shading(cell, '003366')
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)

# 저장
doc.save(OUTPUT_PATH)
print(f'\n=== 변환 완료 ===')
print(f'파일: {OUTPUT_PATH}')

# 임시 파일 정리
import shutil
shutil.rmtree(TEMP_DIR, ignore_errors=True)
